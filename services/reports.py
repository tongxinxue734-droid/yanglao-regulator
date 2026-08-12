# -*- coding: utf-8 -*-
"""报告生成：按养老机构的月度/年度报告（100 分制 · 检查-扣分-报告闭环）"""
import os
from datetime import datetime
from collections import Counter, defaultdict

import pandas as pd
from sqlalchemy.orm import Session

import config
from models import Hazard, ViolationRecord, ComplianceScore, Organization, Indicator
from services.scoring import compliance_score, punishment_tier


def _org(session: Session, org_id: int):
    org = session.query(Organization).get(org_id) if org_id else None
    if org:
        return {"id": org.id, "name": org.name, "code": org.code,
                "base_score": org.base_score or 100}
    return {"id": None, "name": "全辖区汇总", "code": "ALL", "base_score": 100}


# ---------------------------------------------------------------
# 月度报告（按机构）
# ---------------------------------------------------------------
def monthly_report(session: Session, period: str, org_id: int = None) -> dict:
    """period: YYYY-MM；org_id=None 为全辖区汇总"""
    org = _org(session, org_id)
    rows = [h for h in session.query(Hazard).all()
            if h.created_at and h.created_at.strftime("%Y-%m") == period
            and (org_id is None or h.org_id == org_id)]
    prev_period = _prev_month(period)
    prev = [h for h in session.query(Hazard).all()
            if h.created_at and h.created_at.strftime("%Y-%m") == prev_period
            and (org_id is None or h.org_id == org_id)]

    total = len(rows)
    closed = sum(1 for h in rows if h.status in ("closed", "archived"))
    overdued = sum(1 for h in rows if h.overdued)
    rect_rate = round(closed / total * 100, 1) if total else 100.0
    prev_total = len(prev)
    mom = round((total - prev_total) / prev_total * 100, 1) if prev_total else None

    cs = compliance_score(period, session, org_id)

    # 本月违规扣分明细（按指标）
    viols = session.query(ViolationRecord).filter(
        ViolationRecord.period == period,
        ViolationRecord.org_id == org_id if org_id is not None else True).all()
    deduct_detail = Counter()
    deduct_status = Counter()
    for v in viols:
        deduct_detail[v.indicator_code] += v.deducted
        deduct_status[v.status] += 1

    dist = Counter(h.category for h in rows)
    top10 = Counter(h.hazard_type or h.title for h in rows).most_common(10)
    levels = Counter(h.level for h in rows)
    open_items = [h for h in rows if h.status not in ("closed", "archived")]

    return {
        "org": org, "period": period, "total": total, "closed": closed,
        "rect_rate": rect_rate, "overdued": overdued, "mom": mom,
        "compliance": cs, "distribution": dict(dist),
        "top10": top10, "levels": dict(levels),
        "open_items": open_items,
        "deduct_detail": dict(deduct_detail), "deduct_status": dict(deduct_status),
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }


# ---------------------------------------------------------------
# 年度总报告（按机构）
# ---------------------------------------------------------------
def annual_report(session: Session, year: str, org_id: int = None) -> dict:
    org = _org(session, org_id)
    rows = [h for h in session.query(Hazard).all()
            if h.created_at and h.created_at.strftime("%Y") == year
            and (org_id is None or h.org_id == org_id)]
    total = len(rows)
    closed = sum(1 for h in rows if h.status in ("closed", "archived"))

    # 月度趋势（12 个月）
    months = [f"{year}-{m:02d}" for m in range(1, 13)]
    trend = {}
    for m in months:
        n = sum(1 for h in rows if h.created_at.strftime("%Y-%m") == m)
        trend[m[5:]] = n

    # 全年扣分（按机构）
    viols = session.query(ViolationRecord).filter(
        ViolationRecord.period.like(f"{year}-%"),
        ViolationRecord.org_id == org_id if org_id is not None else True).all()
    year_deduct = sum(v.deducted for v in viols)
    tier = punishment_tier(year_deduct)
    # 年度扣分按类别（根因）
    ind_map = {i.code: i for i in session.query(Indicator).all()}
    cat_deduct = Counter()
    for v in viols:
        ind = ind_map.get(v.indicator_code)
        cat_deduct[(ind.category if ind else "其他")] += v.deducted

    # 高频根因（隐患类型 TOP）
    root = Counter(h.hazard_type or h.title for h in rows).most_common(8)
    # 等级分布
    levels = Counter(h.level for h in rows)

    return {
        "org": org, "year": year, "total": total, "closed": closed,
        "rect_rate": round(closed / total * 100, 1) if total else 100.0,
        "trend": trend, "root_causes": root, "levels": dict(levels),
        "cat_deduct": dict(cat_deduct),
        "year_deduct": year_deduct, "tier": tier, "violations": len(viols),
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }


def _prev_month(period: str) -> str:
    y, m = int(period[:4]), int(period[5:7])
    if m == 1:
        return f"{y-1}-12"
    return f"{y}-{m-1:02d}"


# ---------------------------------------------------------------
# 文字分析
# ---------------------------------------------------------------
def analysis_text(rep: dict, rtype: str = "月度") -> str:
    org_name = rep.get("org", {}).get("name", "该机构") if isinstance(rep.get("org"), dict) else "该机构"
    if rtype == "月度":
        cs = rep["compliance"]
        mom = f"较上月{'上升' if (rep['mom'] or 0) > 0 else '下降'} {abs(rep['mom'] or 0)}%" if rep["mom"] is not None else "首月无环比"
        top = "、".join(f"{t}({n}次)" for t, n in rep["top10"][:5]) or "无"
        detail = "；".join(f"{k} 扣 {v} 分" for k, v in list(rep["deduct_detail"].items())[:5]) or "无扣分项"
        return (f"{org_name} 本月共发现隐患 {rep['total']} 项（{mom}），已闭环 {rep['closed']} 项，"
                f"整改完成率 {rep['rect_rate']}%，逾期 {rep['overdued']} 项。"
                f"合规考核累计扣分 {cs['deducted']} 分，综合得分 {cs['score']} 分，"
                f"落入「{cs['tier']['grade']}」档次（{cs['tier']['penalty']}）。"
                f"主要扣分项：{detail}。高频隐患 TOP5：{top}。"
                f"建议：{_monthly_advice(rep)}")
    return (f"{org_name} {rep['year']} 年度共发现隐患 {rep['total']} 项，闭环 {rep['closed']} 项，"
            f"全年整改完成率 {rep['rect_rate']}%，累计合规扣分 {rep['year_deduct']} 分，"
            f"合规档次「{rep['tier']['grade']}」。"
            f"建议：{_annual_advice(rep)}")


def _monthly_advice(rep: dict) -> str:
    tips = []
    if rep["overdued"]:
        tips.append("对逾期隐患逐项督办，落实整改责任人")
    if rep["compliance"]["deducted"]:
        top = list(rep["deduct_detail"].items())[:3]
        tips.append("针对扣分项制定专项整改方案：" + "、".join(f"{k}({v}分)" for k, v in top))
    if rep["total"]:
        cat = max(rep["distribution"].items(), key=lambda x: x[1])[0]
        tips.append(f"重点关注「{cat}」类隐患，开展专项巡检")
    return "；".join(tips) if tips else "保持现有检查与整改机制，持续开展安全培训"


def _annual_advice(rep: dict) -> str:
    tips = []
    if rep["root_causes"]:
        tips.append(f"对高频隐患「{rep['root_causes'][0][0]}」开展根因分析")
    tips.append("完善全员安全生产责任制与应急预案演练")
    if rep["year_deduct"] >= 6:
        tips.append("年度合规评分触发较重档次，建议启动管理层专项整改")
    return "；".join(tips)


# ---------------------------------------------------------------
# 导出：Excel / HTML
# ---------------------------------------------------------------
def export_excel(rep: dict, rtype: str = "月度", out_dir: str = None) -> str:
    out_dir = out_dir or config.REPORT_DIR
    os.makedirs(out_dir, exist_ok=True)
    org = rep.get("org", {})
    org_tag = org.get("code", "ALL") if isinstance(org, dict) else "ALL"
    fname = f"{org_tag}_{rtype}报告_{rep.get('period', rep.get('year'))}.xlsx"
    path = os.path.join(out_dir, fname)
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        if rtype == "月度":
            cs = rep["compliance"]
            pd.DataFrame([{
                "机构": rep["org"]["name"], "期间": rep["period"],
                "隐患总数": rep["total"], "已闭环": rep["closed"],
                "整改完成率%": rep["rect_rate"], "逾期数": rep["overdued"],
                "环比%": rep["mom"], "合规扣分": cs["deducted"],
                "综合得分": cs["score"], "处罚档次": cs["tier"]["grade"],
                "处罚措施": cs["tier"]["penalty"],
            }]).to_excel(w, sheet_name="核心指标", index=False)
            pd.DataFrame([{"类别": k, "数量": v} for k, v in rep["distribution"].items()]
                         ).to_excel(w, sheet_name="隐患分布", index=False)
            pd.DataFrame([{"指标编号": k, "扣分": v} for k, v in rep["deduct_detail"].items()]
                         ).to_excel(w, sheet_name="扣分明细", index=False)
            pd.DataFrame([{"编号": h.code, "标题": h.title, "等级": h.level,
                           "状态": _status_cn(h.status),
                           "上报人": h.reporter.name if h.reporter else "",
                           "整改责任人": h.assignee.name if h.assignee else "",
                           "期限": h.deadline.strftime("%Y-%m-%d") if h.deadline else "",
                           "逾期": "是" if h.overdued else "否"}
                          for h in rep["open_items"]]).to_excel(w, sheet_name="未闭环明细", index=False)
        else:
            pd.DataFrame([{"机构": rep["org"]["name"], "年份": rep["year"],
                           "隐患总数": rep["total"], "闭环数": rep["closed"],
                           "整改完成率%": rep["rect_rate"], "全年扣分": rep["year_deduct"],
                           "合规档次": rep["tier"]["grade"]}]).to_excel(w, sheet_name="年度概览", index=False)
            pd.DataFrame([{"月份": k, "隐患数": v} for k, v in rep["trend"].items()]
                         ).to_excel(w, sheet_name="月度趋势", index=False)
            pd.DataFrame([{"扣分类别": k, "扣分": v} for k, v in rep["cat_deduct"].items()]
                         ).to_excel(w, sheet_name="年度扣分根因", index=False)
    return path


def export_html(rep: dict, rtype: str = "月度") -> str:
    """HTML 报告：带机构抬头、公章位、签字栏，可直接打印为 PDF"""
    org = rep.get("org", {})
    org_name = org.get("name", "全辖区") if isinstance(org, dict) else "全辖区"
    period = rep.get("period", rep.get("year", ""))
    cs = rep.get("compliance", {})
    tier = (cs or {}).get("tier", rep.get("tier", {}))
    if rtype == "月度":
        rows_html = "".join(
            f"<tr><td>{h.code}</td><td>{h.title}</td><td>{h.level}</td>"
            f"<td>{_status_cn(h.status)}</td>"
            f"<td>{h.reporter.name if h.reporter else ''}</td>"
            f"<td>{h.assignee.name if h.assignee else ''}</td>"
            f"<td>{'是' if h.overdued else '否'}</td></tr>"
            for h in rep["open_items"])
        table_head = ("<tr><th>编号</th><th>隐患标题</th><th>等级</th><th>状态</th>"
                      "<th>上报人</th><th>整改责任人</th><th>逾期</th></tr>")
    else:
        rows_html = "".join(f"<tr><td>{m}</td><td>{n}</td></tr>" for m, n in rep["trend"].items())
        table_head = "<tr><th>月份</th><th>隐患数</th></tr>"

    html = f"""<!DOCTYPE html><html lang="zh-CN"><head><meta charset="utf-8">
<style>
body{{font-family:'Microsoft YaHei',sans-serif;margin:40px;color:#333;}}
h1{{text-align:center;font-size:22px;}}
.sub{{text-align:center;color:#666;margin-bottom:24px;}}
table{{border-collapse:collapse;width:100%;margin:12px 0;font-size:13px;}}
th,td{{border:1px solid #999;padding:6px 8px;text-align:left;}}
th{{background:#f0f4f8;}}
.kpi{{display:flex;gap:12px;margin:16px 0;}}
.kpi div{{flex:1;border:1px solid #ccc;border-radius:8px;padding:12px;text-align:center;}}
.kpi b{{display:block;font-size:26px;color:#1f6feb;}}
.sig{{margin-top:48px;display:flex;justify-content:space-between;font-size:14px;}}
.seal{{width:120px;height:120px;border:2px dashed #b00;border-radius:50%;display:flex;
align-items:center;justify-content:center;color:#b00;font-size:12px;text-align:center;}}
</style></head><body>
<h1>{org_name}</h1>
<h1>{rtype}合规检查报告（{period}）</h1>
<div class="sub">报告生成时间：{rep.get('generated_at', '')}</div>
<div style="text-align:center;font-size:12px;color:#888;margin-bottom:18px;border:1px dashed #ccc;padding:6px;">
🤖 本报告含 AI 辅助生成内容（依据 AIGC 内容标识规范标注）· 数据已脱敏处理
</div>
<div class="kpi">
<div><b>{rep.get('total', 0)}</b>隐患总数</div>
<div><b>{rep.get('rect_rate', 0)}%</b>整改完成率</div>
<div><b>{rep.get('overdued', 0)}</b>逾期隐患</div>
<div><b>{cs.get('score', rep.get('year_deduct', 0))}</b>合规得分{'(扣'+str(cs.get('deducted',0))+'分)' if rtype=='月度' else ''}</div>
<div><b>{tier.get('grade', '')}</b>处罚档次</div>
</div>
<h3>一、核心指标概览</h3>
<p>处罚档次对应措施：{tier.get('penalty', '')}</p>
<h3>二、分析结论与改进建议</h3>
<p>{analysis_text(rep, rtype)}</p>
<h3>三、明细（{ '未闭环/逾期隐患' if rtype=='月度' else '月度趋势' }）</h3>
<table>{table_head}{rows_html}</table>
<div class="sig">
<div>检查人：____________　　审核人：____________</div>
<div class="seal">（机构公章）</div>
</div>
</body></html>"""
    return html


def _status_cn(status: str) -> str:
    return config.HAZARD_STATUS.get(status, status)


# ---------------------------------------------------------------
# 公文式 Word 报告导出（python-docx，红头文件格式）
# ---------------------------------------------------------------
def export_docx(rep: dict, rtype: str = "月度", out_dir: str = None) -> str:
    """生成红头公文式 Word 报告：标题(红) + 文号 + 正文 + 落款 + 签章位"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    out_dir = out_dir or config.REPORT_DIR
    os.makedirs(out_dir, exist_ok=True)
    org = rep.get("org", {})
    org_name = org.get("name", "全辖区") if isinstance(org, dict) else "全辖区"
    org_code = org.get("code", "ALL") if isinstance(org, dict) else "ALL"
    period = rep.get("period", rep.get("year", ""))
    fname = f"{org_code}_{rtype}检查报告_{period}.docx"
    path = os.path.join(out_dir, fname)

    doc = Document()
    # 页面边距（公文标准）
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

    def set_font(run, name="仿宋_GB2312", size=14, bold=False, color=None):
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    # ---- 红头标题 ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("西安市养老机构运营监管检查报告")
    set_font(r, name="方正小标宋简体", size=22, bold=True, color=(0xB0, 0x00, 0x00))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"——{org_name}（{org_code}）{rtype}报告（{period}）")
    set_font(r2, name="方正小标宋简体", size=16, bold=True, color=(0xB0, 0x00, 0x00))
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("（此件依 AIGC 内容标识规范标注，含 AI 辅助生成内容）")
    set_font(r3, name="楷体_GB2312", size=10, color=(0x66, 0x66, 0x66))

    # ---- 文号 ----
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(f"市民函〔{period[:4]}〕第 {org_code.split('-')[-1]} 号")
    set_font(r4, name="仿宋_GB2312", size=12)

    doc.add_paragraph()

    # ---- 正文 ----
    def para(text, indent=True, bold=False):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        r = p.add_run(text)
        set_font(r, bold=bold)
        return p

    cs = rep.get("compliance", {})
    tier = (cs or {}).get("tier", rep.get("tier", {}))
    para(f"一、检查概况", bold=True)
    if rtype == "月度":
        para(f"本月共对该机构实施检查，发现问题隐患 {rep.get('total', 0)} 项，已闭环 "
             f"{rep.get('closed', 0)} 项，整改完成率 {rep.get('rect_rate', 0)}%，"
             f"逾期未整改 {rep.get('overdued', 0)} 项。")
        para(f"依据《养老机构运营违规评价指标》，本月累计扣分 {cs.get('deducted', 0)} 分，"
             f"机构综合得分 {cs.get('score', 0)} 分，落入「{tier.get('grade', '')}」处罚档次，"
             f"对应处置措施：{tier.get('penalty', '')}。")
    else:
        para(f"年度共发现隐患 {rep.get('total', 0)} 项，闭环 {rep.get('closed', 0)} 项，"
             f"全年整改完成率 {rep.get('rect_rate', 0)}%，全年累计合规扣分 "
             f"{rep.get('year_deduct', 0)} 分，合规档次「{tier.get('grade', '')}」。")

    para("二、主要扣分事项", bold=True)
    detail = rep.get("deduct_detail") if rtype == "月度" else rep.get("cat_deduct")
    if detail:
        for code, ded in list(detail.items())[:8]:
            para(f"（一）{code}：扣 {ded} 分。", indent=True)
    else:
        para("本期无违规扣分事项。")

    para("三、分析结论与整改要求", bold=True)
    para(analysis_text(rep, rtype))

    doc.add_paragraph()
    # ---- 落款 ----
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r5 = p5.add_run("西安市民政局养老服务科")
    set_font(r5, name="仿宋_GB2312", size=14)
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r6 = p6.add_run(f"{rep.get('generated_at', '')[:10]}")
    set_font(r6, name="仿宋_GB2312", size=14)
    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r7 = p7.add_run("（盖章）")
    set_font(r7, name="仿宋_GB2312", size=12, color=(0xB0, 0x00, 0x00))

    doc.save(path)
    return path
